from __future__ import annotations

from collections import Counter
from pathlib import Path
import re

import fitz
from docx import Document


class ExtractionError(Exception):
    pass


def detect_file_type(file_path: str) -> str:
    with open(file_path, "rb") as handle:
        magic = handle.read(4)
    if magic == b"%PDF":
        return "pdf"
    if magic == b"PK\x03\x04":
        return "docx"
    raise ExtractionError("Unsupported file type")


def _normalize_lines(lines: list[str]) -> str:
    cleaned: list[str] = []
    blank = False
    for line in lines:
        stripped = re.sub(r"\s+", " ", line).strip()
        if not stripped:
            if not blank:
                cleaned.append("")
            blank = True
            continue
        blank = False
        if re.fullmatch(r"\d+|page \d+( of \d+)?", stripped, flags=re.IGNORECASE):
            continue
        cleaned.append(stripped)
    return "\n".join(cleaned).strip()


def extract_text_from_pdf(file_path: str) -> str:
    try:
        doc = fitz.open(file_path)
    except Exception as exc:
        raise ExtractionError("PDF is corrupt or unreadable") from exc

    if doc.is_encrypted:
        raise ExtractionError("PDF is encrypted")

    page_lines: list[list[str]] = []
    try:
        for page in doc:
            blocks = page.get_text("blocks")
            blocks = sorted(blocks, key=lambda block: (round(block[1] / 10) * 10, block[0]))
            lines = []
            for block in blocks:
                text = str(block[4]).strip()
                if text:
                    lines.extend(text.splitlines())
            page_lines.append(lines)
    finally:
        doc.close()

    repeated_candidates = Counter(
        line.strip()
        for lines in page_lines
        for line in lines
        if 0 < len(line.strip().split()) < 4
    )
    repeated = {line for line, count in repeated_candidates.items() if count > 1}
    clean_lines = [line for lines in page_lines for line in lines if line.strip() not in repeated]
    text = _normalize_lines(clean_lines)
    if len(text) < 100:
        raise ExtractionError("Extracted text is too short")
    return text


def extract_text_from_docx(file_path: str) -> str:
    try:
        document = Document(file_path)
    except Exception as exc:
        raise ExtractionError("DOCX is corrupt or unreadable") from exc

    lines: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs)

    text = _normalize_lines(lines)
    if len(text) < 100:
        raise ExtractionError("Extracted text is too short")
    return text


def extract_text(file_path: str) -> str:
    file_type = detect_file_type(file_path)
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    if file_type == "docx":
        return extract_text_from_docx(file_path)
    raise ExtractionError(f"Unsupported file type for {Path(file_path).name}")
