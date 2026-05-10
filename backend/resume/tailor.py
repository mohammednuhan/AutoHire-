from __future__ import annotations

import html
import json
import logging
import re
from copy import deepcopy
from pathlib import Path
from typing import Any
from uuid import uuid4

from pydantic import BaseModel, ValidationError

from database.models import Application, Job
from database.session import AsyncSessionLocal
from llm.client import LLMFailure, LLMRouter
from llm.prompts import RESUME_TAILORING_PROMPT, RESUME_TAILORING_SYSTEM
from resume.formatting import extract_key_phrases, flatten_profile_terms, format_full_profile
from resume.safety import ContentSafetyError, sanitize_user_content
from schemas.api_schemas import ExperienceItem, ProjectItem, ResumeProfile, SkillsProfile
from storage import data_dir

logger = logging.getLogger("autohire.resume.tailor")

DATE_SEPARATOR = " \u2013 "


class TailoredResumeResult(BaseModel):
    success: bool
    pdf_path: str | None = None
    docx_path: str | None = None
    profile: ResumeProfile | None = None
    used_fallback: bool = False
    failure_reason: str | None = None


class TailoringError(Exception):
    pass


async def tailor_resume(
    profile: ResumeProfile,
    job: Job,
    application_id: str,
    llm_router: LLMRouter,
) -> TailoredResumeResult:
    used_fallback = False
    try:
        tailored = await _tailor_with_llm(profile, job, llm_router)
        _assert_traceable(profile, tailored, job)
    except (TailoringError, ValidationError, ValueError) as exc:
        logger.warning(
            "resume_tailoring_llm_failed",
            extra={"application_id": application_id, "job_id": job.id, "error": str(exc)},
        )
        tailored = _deterministic_tailor(profile, job)
        used_fallback = True

    try:
        output_dir = data_dir() / "applications" / application_id
        output_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = output_dir / "resume_tailored.pdf"
        docx_path = output_dir / "resume_tailored.docx"
        _write_pdf(tailored, pdf_path)
        _write_docx(tailored, docx_path)
        await _update_application_paths(application_id, pdf_path, docx_path)
    except Exception as exc:
        logger.exception("resume_tailoring_output_failed", extra={"application_id": application_id})
        return TailoredResumeResult(success=False, failure_reason=str(exc))

    return TailoredResumeResult(
        success=True,
        pdf_path=str(pdf_path),
        docx_path=str(docx_path),
        profile=tailored,
        used_fallback=used_fallback,
    )


async def _tailor_with_llm(profile: ResumeProfile, job: Job, llm_router: LLMRouter) -> ResumeProfile:
    try:
        safe_title = sanitize_user_content(job.title)[:500]
        safe_description = sanitize_user_content(job.description or "")[:4000]
    except ContentSafetyError as exc:
        raise TailoringError(f"Unsafe job content rejected: {exc}") from exc

    prompt = RESUME_TAILORING_PROMPT.format(
        job_title=safe_title,
        job_skills_required=job.skills_required or [],
        key_phrases=extract_key_phrases(safe_description),
        full_profile=format_full_profile(profile),
    )
    try:
        response = await llm_router.call_with_retry(
            task_type="write",
            prompt=prompt,
            system=RESUME_TAILORING_SYSTEM,
            response_format="json",
            max_retries=3,
            trace_id=str(uuid4()),
        )
    except LLMFailure as exc:
        raise TailoringError(str(exc)) from exc

    return ResumeProfile.model_validate(_json_payload(response))


def _deterministic_tailor(profile: ResumeProfile, job: Job) -> ResumeProfile:
    data = deepcopy(profile.model_dump())
    tailored = ResumeProfile.model_validate(data)

    phrases = set(_normalize(value) for value in (job.skills_required or []))
    phrases.update(_normalize(value) for value in extract_key_phrases(job.description))
    phrases = {phrase for phrase in phrases if phrase}

    tailored.skills = _reorder_skills(profile.skills, phrases)
    tailored.experience = _rank_experience(profile.experience, phrases)
    tailored.projects = _rank_projects(profile.projects, phrases)
    tailored.summary = _summary_for_job(tailored, job, tailored.skills)
    return tailored


def _reorder_skills(skills: SkillsProfile, phrases: set[str]) -> SkillsProfile:
    reordered: dict[str, list[str]] = {}
    for category, values in skills.model_dump().items():
        reordered[category] = sorted(
            values,
            key=lambda value: (0 if _matches(value, phrases) else 1, str(value).lower()),
        )
    return SkillsProfile.model_validate(reordered)


def _rank_experience(experience: list[ExperienceItem], phrases: set[str]) -> list[ExperienceItem]:
    ranked = sorted(experience, key=lambda item: _experience_score(item, phrases), reverse=True)
    if any(_experience_score(item, phrases) > 0 for item in ranked):
        return [item for item in ranked if _experience_score(item, phrases) > 0]
    return ranked


def _rank_projects(projects: list[ProjectItem], phrases: set[str]) -> list[ProjectItem]:
    ranked = sorted(projects, key=lambda item: _project_score(item, phrases), reverse=True)
    if any(_project_score(item, phrases) > 0 for item in ranked):
        return [item for item in ranked if _project_score(item, phrases) > 0]
    return ranked


def _summary_for_job(profile: ResumeProfile, job: Job, skills: SkillsProfile) -> str:
    skill_pool = _all_skills(skills)
    selected_skills = ", ".join(skill_pool[:5]) if skill_pool else "the documented technical background"
    first_sentence = f"{profile.full_name} is targeting the {job.title} role with experience across {selected_skills}."
    evidence = _best_evidence(profile)
    second_sentence = f"This version emphasizes traceable resume evidence from {evidence} for {job.company}."
    return f"{first_sentence} {second_sentence}"


def _best_evidence(profile: ResumeProfile) -> str:
    names = [project.name for project in profile.projects[:2]]
    names.extend(f"{item.role} at {item.company}" for item in profile.experience[:1])
    return ", ".join(names) if names else "the original profile"


def _assert_traceable(original: ResumeProfile, tailored: ResumeProfile, job: Job) -> None:
    if tailored.full_name != original.full_name:
        raise ValueError("full_name cannot be changed")
    for field in ("email", "phone", "location", "linkedin_url", "github_url", "portfolio_url"):
        if getattr(tailored, field) != getattr(original, field):
            raise ValueError(f"{field} cannot be changed")

    original_skills = {
        category: {_normalize(value) for value in values}
        for category, values in original.skills.model_dump().items()
    }
    for category, values in tailored.skills.model_dump().items():
        new_values = {_normalize(value) for value in values} - original_skills.get(category, set())
        if new_values:
            raise ValueError(f"New skills are not traceable: {sorted(new_values)}")

    original_experience = {
        (item.company, item.role, item.start_date, item.end_date): item for item in original.experience
    }
    for item in tailored.experience:
        key = (item.company, item.role, item.start_date, item.end_date)
        if key not in original_experience:
            raise ValueError(f"Experience changed or added: {item.role} at {item.company}")
        original_tech = {_normalize(value) for value in original_experience[key].tech_stack}
        if {_normalize(value) for value in item.tech_stack} - original_tech:
            raise ValueError(f"Experience tech_stack added untraceable values: {item.role}")

    original_projects = {item.name: item for item in original.projects}
    for item in tailored.projects:
        if item.name not in original_projects:
            raise ValueError(f"Project changed or added: {item.name}")
        original_tech = {_normalize(value) for value in original_projects[item.name].tech_stack}
        if {_normalize(value) for value in item.tech_stack} - original_tech:
            raise ValueError(f"Project tech_stack added untraceable values: {item.name}")

    if _education_keys(tailored) != _education_keys(original):
        raise ValueError("Education entries cannot be changed")
    if _certification_keys(tailored) != _certification_keys(original):
        raise ValueError("Certification entries cannot be changed")
    if set(tailored.achievements) != set(original.achievements):
        raise ValueError("Achievements cannot be changed")
    if set(tailored.languages_spoken) != set(original.languages_spoken):
        raise ValueError("Languages spoken cannot be changed")

    original_terms = flatten_profile_terms(original)
    tailored_terms = flatten_profile_terms(tailored)
    untraceable_terms = tailored_terms - original_terms
    allowed_terms = {_normalize(tailored.full_name)}
    if untraceable_terms - allowed_terms:
        examples = sorted(untraceable_terms - allowed_terms)[:5]
        raise ValueError(f"Tailored profile includes untraceable terms: {examples}")

    _assert_no_new_numbers(original, tailored, job)
    _assert_no_unsupported_required_skills(original, tailored, job)


def _assert_no_new_numbers(original: ResumeProfile, tailored: ResumeProfile, job: Job) -> None:
    original_numbers = _number_tokens(original.model_dump())
    allowed_job_numbers = _number_tokens(
        {
            "title": job.title,
            "company": job.company,
            "skills_required": job.skills_required or [],
        }
    )
    new_numbers = _number_tokens(tailored.model_dump()) - original_numbers - allowed_job_numbers
    if new_numbers:
        raise ValueError(f"Tailored profile includes untraceable numbers: {sorted(new_numbers)}")


def _assert_no_unsupported_required_skills(
    original: ResumeProfile,
    tailored: ResumeProfile,
    job: Job,
) -> None:
    original_text = _normalize(json.dumps(original.model_dump(), ensure_ascii=True))
    tailored_text = _normalize(json.dumps(tailored.model_dump(), ensure_ascii=True))
    unsupported_mentions: list[str] = []
    for skill in job.skills_required or []:
        normalized_skill = _normalize(skill)
        if (
            normalized_skill
            and normalized_skill not in original_text
            and normalized_skill in tailored_text
        ):
            unsupported_mentions.append(skill)
    if unsupported_mentions:
        raise ValueError(
            "Tailored profile mentions required skills not present in the original profile: "
            f"{sorted(set(unsupported_mentions))}"
        )


def _number_tokens(payload: Any) -> set[str]:
    return set(re.findall(r"\d+(?:[.,]\d+)?%?", json.dumps(payload, ensure_ascii=True)))


def _write_pdf(profile: ResumeProfile, output_path: Path) -> None:
    from weasyprint import HTML

    HTML(string=_profile_to_html(profile), base_url=str(output_path.parent)).write_pdf(str(output_path))


def _education_keys(profile: ResumeProfile) -> set[tuple[Any, ...]]:
    return {
        (
            item.institution,
            item.degree,
            item.field,
            item.graduation_year,
            item.gpa,
            tuple(item.relevant_courses),
        )
        for item in profile.education
    }


def _certification_keys(profile: ResumeProfile) -> set[tuple[Any, ...]]:
    return {
        (
            item.name,
            item.issuer,
            item.year,
        )
        for item in profile.certifications
    }


def _write_docx(profile: ResumeProfile, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title_run = title.add_run(profile.full_name)
    title_run.bold = True
    title_run.font.size = Pt(14)
    _add_contact_line(document, profile)
    _add_docx_section(document, "Summary", [profile.summary] if profile.summary else [])
    _add_docx_experience(document, profile.experience)
    _add_docx_projects(document, profile.projects)
    _add_docx_education(document, profile.education)
    _add_docx_skills(document, profile.skills)
    _add_docx_section(document, "Certifications", [_certification_line(item) for item in profile.certifications])
    document.save(str(output_path))


def _profile_to_html(profile: ResumeProfile) -> str:
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    @page {{ size: A4; margin: 0.55in; }}
    body {{ font-family: Arial, Calibri, sans-serif; font-size: 10.5pt; line-height: 1.25; color: #111; }}
    h1 {{ font-size: 16pt; margin: 0 0 3px; }}
    h2 {{ font-size: 11.5pt; margin: 12px 0 4px; border-bottom: 1px solid #222; }}
    p {{ margin: 0 0 4px; }}
    ul {{ margin: 2px 0 6px 18px; padding: 0; }}
    li {{ margin: 0 0 2px; }}
    .contact {{ margin-bottom: 8px; }}
    .item {{ margin-bottom: 7px; }}
    .item-title {{ font-weight: bold; }}
    .date {{ float: right; font-weight: normal; }}
  </style>
</head>
<body>
  <h1>{_h(profile.full_name)}</h1>
  <p class="contact">{_h(_contact_line(profile))}</p>
  {_html_section("Summary", [_h(profile.summary)] if profile.summary else [])}
  {_experience_html(profile.experience)}
  {_projects_html(profile.projects)}
  {_education_html(profile.education)}
  {_skills_html(profile.skills)}
  {_html_section("Certifications", [_h(_certification_line(item)) for item in profile.certifications])}
</body>
</html>"""


def _experience_html(experience: list[ExperienceItem]) -> str:
    items = []
    for item in experience:
        date = _date_range(item.start_date, item.end_date, item.is_current)
        bullets = "".join(f"<li>{_h(line)}</li>" for line in item.description)
        tech = f"<p><strong>Tech:</strong> {_h(', '.join(item.tech_stack))}</p>" if item.tech_stack else ""
        items.append(
            f"""<div class="item">
  <p class="item-title">{_h(item.role)} - {_h(item.company)}<span class="date">{_h(date)}</span></p>
  {tech}
  <ul>{bullets}</ul>
</div>"""
        )
    return _html_block("Experience", items)


def _projects_html(projects: list[ProjectItem]) -> str:
    items = []
    for item in projects:
        tech = f" | {_h(', '.join(item.tech_stack))}" if item.tech_stack else ""
        duration = f"<span class=\"date\">{_h(item.duration or '')}</span>" if item.duration else ""
        description = f"<p>{_h(item.description)}</p>" if item.description else ""
        items.append(
            f"""<div class="item">
  <p class="item-title">{_h(item.name)}{duration}</p>
  {description}
  <p>{tech.lstrip(' | ')}</p>
</div>"""
        )
    return _html_block("Projects", items)


def _education_html(education: list[Any]) -> str:
    lines = []
    for item in education:
        detail = ", ".join(part for part in [item.degree, item.field] if part)
        year = str(item.graduation_year) if item.graduation_year else ""
        lines.append(
            f"<p><strong>{_h(item.institution)}</strong> - {_h(detail)} "
            f"<span class=\"date\">{_h(year)}</span></p>"
        )
    return _html_block("Education", lines)


def _skills_html(skills: SkillsProfile) -> str:
    lines = []
    labels = {
        "languages": "Languages",
        "frameworks": "Frameworks",
        "databases": "Databases",
        "tools": "Tools",
        "cloud": "Cloud",
        "soft_skills": "Soft Skills",
    }
    for key, values in skills.model_dump().items():
        if values:
            lines.append(f"<p><strong>{labels.get(key, key.title())}:</strong> {_h(', '.join(values))}</p>")
    return _html_block("Skills", lines)


def _html_section(title: str, lines: list[str]) -> str:
    if not lines:
        return ""
    return _html_block(title, [f"<p>{line}</p>" for line in lines])


def _html_block(title: str, body: list[str]) -> str:
    if not body:
        return ""
    return f"<h2>{_h(title)}</h2>\n" + "\n".join(body)


def _add_contact_line(document: Any, profile: ResumeProfile) -> None:
    line = _contact_line(profile)
    if line:
        document.add_paragraph(line)


def _add_docx_section(document: Any, title: str, lines: list[str]) -> None:
    lines = [line for line in lines if line]
    if not lines:
        return
    _add_heading(document, title)
    for line in lines:
        document.add_paragraph(line)


def _add_docx_experience(document: Any, experience: list[ExperienceItem]) -> None:
    if not experience:
        return
    _add_heading(document, "Experience")
    for item in experience:
        date_range = _date_range(item.start_date, item.end_date, item.is_current)
        document.add_paragraph(f"{item.role} - {item.company} | {date_range}")
        if item.tech_stack:
            document.add_paragraph(f"Tech: {', '.join(item.tech_stack)}")
        for bullet in item.description:
            document.add_paragraph(bullet, style="List Bullet")


def _add_docx_projects(document: Any, projects: list[ProjectItem]) -> None:
    if not projects:
        return
    _add_heading(document, "Projects")
    for item in projects:
        suffix = f" | {item.duration}" if item.duration else ""
        document.add_paragraph(f"{item.name}{suffix}")
        if item.description:
            document.add_paragraph(item.description)
        if item.tech_stack:
            document.add_paragraph(", ".join(item.tech_stack))


def _add_docx_education(document: Any, education: list[Any]) -> None:
    if not education:
        return
    _add_heading(document, "Education")
    for item in education:
        detail = ", ".join(part for part in [item.degree, item.field] if part)
        year = f" | {item.graduation_year}" if item.graduation_year else ""
        document.add_paragraph(f"{item.institution} - {detail}{year}")


def _add_docx_skills(document: Any, skills: SkillsProfile) -> None:
    labels = {
        "languages": "Languages",
        "frameworks": "Frameworks",
        "databases": "Databases",
        "tools": "Tools",
        "cloud": "Cloud",
        "soft_skills": "Soft Skills",
    }
    lines = [
        f"{labels.get(key, key.title())}: {', '.join(values)}"
        for key, values in skills.model_dump().items()
        if values
    ]
    _add_docx_section(document, "Skills", lines)


def _add_heading(document: Any, title: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = document.styles["Normal"].font.size


def _contact_line(profile: ResumeProfile) -> str:
    parts = [
        profile.email,
        profile.phone,
        profile.location,
        profile.linkedin_url,
        profile.github_url,
        profile.portfolio_url,
    ]
    return " | ".join(part for part in parts if part)


def _certification_line(item: Any) -> str:
    parts = [item.name, item.issuer, str(item.year) if item.year else None]
    return " | ".join(part for part in parts if part)


def _date_range(start_date: str | None, end_date: str | None, is_current: bool = False) -> str:
    start = _format_date(start_date)
    end = "Present" if is_current else _format_date(end_date)
    if start and end:
        return f"{start}{DATE_SEPARATOR}{end}"
    return start or end


def _format_date(value: str | None) -> str:
    if not value:
        return ""
    text = str(value).strip()
    if text.lower() == "present":
        return "Present"
    match = re.fullmatch(r"(\d{4})-(\d{1,2})", text)
    if match:
        year, month = match.groups()
        month_names = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        month_index = max(1, min(12, int(month))) - 1
        return f"{month_names[month_index]} {year}"
    return text


def _all_skills(skills: SkillsProfile) -> list[str]:
    values: list[str] = []
    for items in skills.model_dump().values():
        values.extend(items)
    return values


def _experience_score(item: ExperienceItem, phrases: set[str]) -> int:
    text = " ".join([item.company, item.role, " ".join(item.description), " ".join(item.tech_stack)])
    return _score_text(text, phrases)


def _project_score(item: ProjectItem, phrases: set[str]) -> int:
    text = " ".join([item.name, item.description or "", " ".join(item.tech_stack)])
    return _score_text(text, phrases)


def _score_text(text: str, phrases: set[str]) -> int:
    normalized = _normalize(text)
    return sum(1 for phrase in phrases if phrase and phrase in normalized)


def _matches(value: str, phrases: set[str]) -> bool:
    normalized = _normalize(value)
    return any(normalized == phrase or normalized in phrase or phrase in normalized for phrase in phrases)


def _normalize(value: str | None) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def _h(value: str | None) -> str:
    return html.escape(value or "")


def _json_payload(response: str) -> dict[str, Any]:
    text = response.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:].strip()
    return json.loads(text)


async def _update_application_paths(application_id: str, pdf_path: Path, docx_path: Path) -> None:
    async with AsyncSessionLocal() as db:
        application = await db.get(Application, application_id)
        if application is None:
            raise TailoringError(f"Application not found: {application_id}")
        application.tailored_resume_pdf_path = str(pdf_path)
        application.tailored_resume_docx_path = str(docx_path)
        await db.commit()
